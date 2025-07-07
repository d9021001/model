# coding: utf-8
"""Insert a red §4.1.2 'Interpretability of the TPX pipeline' subsection after §4.1.1,
before §4.2, in the REVISED paper-1 docx. Preserves all existing content (incl. OMML)."""
import sys, warnings, copy
warnings.filterwarnings("ignore"); sys.stdout.reconfigure(encoding="utf-8")
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn

F = "Manuscript-app-addict-Jba2026_0101a_REVISED.docx"
d = docx.Document(F)
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath"
before_math = len(d.element.body.findall(".//" + M)); before_paras = len(d.paragraphs)

# locate the §4.2 heading paragraph (insert our subsection before it)
paras = list(d.paragraphs)
hi = next(i for i, p in enumerate(paras) if p.text.strip().startswith("4.1.1"))
heading_src = paras[hi]; body_src = paras[hi + 1]
target = next(p for p in paras if p.text.strip().startswith("4.2 Comparison"))

# capture formatting from existing §4.1.1 heading + body runs
hr = heading_src.runs[0]; br = body_src.runs[0]
h_name, h_size, h_bold = hr.font.name, hr.font.size, hr.font.bold
b_name, b_size = br.font.name, br.font.size
RED = RGBColor(0xFF, 0x00, 0x00)

def styled(run, name, size, bold):
    run.font.color.rgb = RED
    if name: run.font.name = name
    if size: run.font.size = size
    run.font.bold = bool(bold)
    # keep east-asian font binding consistent if present
    if name:
        rPr = run._element.get_or_add_rPr(); rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {}); rPr.insert(0, rFonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs"): rFonts.set(qn(a), name)

HEAD = "4.1.2 Interpretability of the TPX pipeline"
BODY = ("Beyond predictive accuracy, the TPX pipeline is deliberately staged to remain partially "
        "interpretable rather than opaque. Its inputs are themselves meaningful behavioural "
        "quantities—per-app weekly usage time—so the model reasons over clinically legible "
        "signals rather than abstract sensor streams. The principal-component stage compresses the "
        "learned triplet embedding into three orthogonal components that retain roughly 84% of the "
        "variance, yielding a compact, low-dimensional representation that can be visualised in three "
        "dimensions and examined for class separation; the XGBoost stage then exposes the relative "
        "contribution of these components through its split-based importance scores, so a risk score "
        "can be traced to its dominant factors rather than accepted on faith (Chen & Guestrin, 2016). "
        "Because the embedding is learned from the usage tensors, these components are combinations of "
        "app-usage features rather than single named applications; full feature-level attribution over "
        "individual apps—for example, SHAP values computed directly on the original usage "
        "variables—is a natural extension that we pursue in companion work using an explicitly "
        "interpretable feature model. The staged design thus trades a degree of end-to-end opacity for "
        "a transparent, auditable decision pathway suited to mental-health screening.")

h = target.insert_paragraph_before("", style=heading_src.style)
styled(h.add_run(HEAD), h_name, h_size, True)
b = target.insert_paragraph_before("", style=body_src.style)
styled(b.add_run(BODY), b_name, b_size, False)

d.save(F)
d2 = docx.Document(F)
after_math = len(d2.element.body.findall(".//" + M)); after_paras = len(d2.paragraphs)
print(f"paras {before_paras} -> {after_paras} (+{after_paras-before_paras});  oMath {before_math} -> {after_math}")
print("inserted heading:", any(p.text.strip()==HEAD for p in d2.paragraphs))
