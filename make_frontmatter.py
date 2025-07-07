# coding: utf-8
"""Generate title pages + cover letters for paper 1 (TPX) and paper 2 (multi-signal),
targeted at Journal of Behavioral Addictions. Author-specific fields are [bracketed]."""
import sys, warnings
warnings.filterwarnings("ignore"); sys.stdout.reconfigure(encoding="utf-8")
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

JOURNAL = "Journal of Behavioral Addictions"
DATE = "19 June 2026"
IRB = ("The study protocol was reviewed and approved by the Institutional Review Board of China "
       "Medical University Hospital, Taichung, Taiwan (IRB No.: CRREC-108-001). All participants "
       "provided written informed consent prior to enrollment, after being fully informed of the "
       "study's aims, procedures, data-confidentiality measures, and their right to withdraw at any "
       "time without penalty. Passive data logging via Android's “App Usage” service was "
       "conducted only after explicit consent for digital-activity tracking and secure data handling.")
AFFIL = ("¹ [Department / Graduate Institute], China Medical University, Taichung, Taiwan\n"
         "² [Department, Affiliated Institution, City, Country]")
CORR = ("* Correspondence: [Corresponding Author Name], [Department / Graduate Institute], China "
        "Medical University, [No. 91, Hsueh-Shih Road, North District], Taichung 404, Taiwan. "
        "E-mail: [m8951016@gmail.com]. Tel.: [+886-x-xxxx-xxxx]. ORCID: [0000-0000-0000-0000].")
AUTHORS = "[First Author]¹, [Second Author]², […], [Corresponding Author]¹,*"

def newdoc():
    d = docx.Document(); f = d.styles["Normal"].font; f.name = "Times New Roman"; f.size = Pt(12)
    return d
def P(d, t="", bold=False, italic=False, center=False, size=None, after=6, before=0):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if t:
        r = p.add_run(t); r.bold = bold; r.italic = italic
        if size: r.font.size = Pt(size)
    return p
def field(d, label, value):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + ": "); r.bold = True
    p.add_run(value); return p
def section(d, head, body):
    P(d, head, bold=True, after=2, before=8)
    P(d, body, after=6)

# ============================ TITLE PAGES ============================
def title_page(out, title, running, keywords, art_type, counts, datavail, paper_no):
    d = newdoc()
    P(d, "Title Page", bold=True, center=True, size=13, after=10)
    field(d, "Article type", art_type)
    field(d, "Submitted to", JOURNAL)
    P(d)
    P(d, title, bold=True, center=True, size=14, after=8)
    P(d, "Running head: " + running, italic=True, center=True, after=10)
    P(d, "Authors", bold=True, after=2)
    P(d, AUTHORS, after=6)
    P(d, "Affiliations", bold=True, after=2)
    for line in AFFIL.split("\n"): P(d, line, after=2)
    P(d, "", after=4)
    P(d, "Corresponding author", bold=True, after=2)
    P(d, CORR, after=8)
    field(d, "Keywords", keywords)
    field(d, "Word count (abstract)", counts["abs"])
    field(d, "Word count (main text, excl. abstract, references, tables/figure legends)", counts["main"])
    field(d, "Number of figures", counts["fig"]); field(d, "Number of tables", counts["tab"])
    field(d, "Number of references", counts["ref"])
    d.add_paragraph()
    section(d, "Funding", "This research received no specific grant from any funding agency in the "
            "public, commercial, or not-for-profit sectors. [If applicable, replace with: This work was "
            "supported by [Funder] under Grant No. [number].]")
    section(d, "Conflict of interest", "The authors declare that they have no conflict of interest.")
    section(d, "Ethics approval and informed consent", IRB)
    section(d, "Data availability", datavail)
    section(d, "Authors' contributions",
            "[Initials] conceived and designed the study; [Initials] developed the software and performed "
            "the analyses; [Initials] curated the data; [Initials] interpreted the results; [Initials] drafted "
            "the manuscript; [Initials] critically revised it; [Initials] supervised the work. All authors "
            "read and approved the final manuscript and agree to be accountable for all aspects of the work. "
            "(Please complete using CRediT taxonomy roles.)")
    section(d, "Acknowledgements", "[Optional: acknowledge non-author contributors, technical support, "
            "or participating institutions. If none, state “Not applicable.”]")
    d.save(out); print("saved", out)

# Paper 1
title_page(
    "Paper1_TitlePage.docx",
    "Detecting High-Risk Smartphone Addiction among College Students via AI-Based App Usage Analysis During COVID-19",
    "AI Detection of Smartphone Addiction",
    "Smartphone addiction; App usage; Android apps; Convolutional neural network; Machine learning; Digital phenotyping",
    "Full-length report",
    dict(abs="205 words", main="≈ 7,460 words", fig="3", tab="3", ref="24"),
    "The source code implementing the Triplet-PCA-XGBoost pipeline and the derived feature matrices that "
    "support the findings are available from the corresponding author upon reasonable request. Raw passive "
    "app-usage logs cannot be shared publicly to protect participant privacy and in accordance with the "
    "ethics approval.",
    1)

# Paper 2
title_page(
    "Paper2_TitlePage.docx",
    "Interpretable Personalized Prediction of Smartphone-Addiction Risk from App Usage Time and Frequency: "
    "A Multi-Signal Ensemble with Green-Learning Features and SHAP Analysis",
    "Personalized App-Usage Prediction of Smartphone Addiction",
    "Smartphone addiction; App usage; Usage frequency; Green learning; SHAP; Personalized prediction; Machine learning",
    "Full-length report",
    dict(abs="304 words [trim to ≤ 250 for submission]", main="≈ 2,570 words", fig="3", tab="5", ref="13"),
    "A self-contained, de-identified analysis bundle—comprising the tidy participant-week-app dataset and "
    "all feature-engineering, modelling, design-search, interpretability (SHAP) and ablation code, with a "
    "script that regenerates the dataset from the source logs—is available from the corresponding author "
    "upon reasonable request. Raw device logs cannot be shared publicly to protect participant privacy.",
    2)

# ============================ COVER LETTERS ============================
def cover_letter(out, paras, sign_name="[Corresponding Author]"):
    d = newdoc()
    P(d, DATE, after=10)
    P(d, "The Editor-in-Chief", after=0); P(d, JOURNAL, after=10)
    P(d, "Dear Editor-in-Chief,", after=8)
    for pr in paras: P(d, pr, after=8)
    P(d, "Sincerely,", after=2, before=6)
    P(d, sign_name + ", on behalf of all co-authors", after=0)
    P(d, "China Medical University, Taichung, Taiwan", after=0)
    P(d, "E-mail: [m8951016@gmail.com]", after=0)
    d.save(out); print("saved", out)

P1 = [
 "Please find enclosed our manuscript entitled “Detecting High-Risk Smartphone Addiction among "
 "College Students via AI-Based App Usage Analysis During COVID-19,” which we respectfully submit "
 "for consideration as a Full-length report in the " + JOURNAL + ".",
 "Problematic smartphone use is typically screened with self-report instruments that are burdensome "
 "and susceptible to bias. Using 24 weeks of passively logged Android app-usage data from 114 "
 "undergraduates observed during the COVID-19 period, we developed an interpretable Triplet-PCA-XGBoost "
 "(TPX) pipeline that detects high-risk individuals from app-usage time. In five-fold cross-validation the "
 "model attained an ROC-AUC of 0.981 and a precision-recall AUC of 0.932 (0.932 ± 0.006 across folds), "
 "with explicit handling of class imbalance and a transparent, inspectable decision pathway.",
 "We believe these findings suit the journal's readership because they (i) demonstrate that an objective "
 "behavioural signal can complement questionnaire-based screening, (ii) ground the predictive validity of "
 "usage time in the addiction construct itself, and (iii) are reported transparently, including the "
 "distinction between within-cohort (baseline-anchored) performance and the more demanding cold-start "
 "setting for previously unseen individuals.",
 "In the interest of full transparency, a companion manuscript (“Interpretable Personalized Prediction "
 "of Smartphone-Addiction Risk from App Usage Time and Frequency …”) drawing on the same cohort is "
 "being submitted separately. The two papers are complementary and non-overlapping: the present paper "
 "introduces the time-only TPX detection pipeline with a fixed per-participant label, whereas the companion "
 "paper studies a multi-signal (time + frequency + category + app identity) interpretable ensemble for "
 "personalized, within-subject prediction of a time-varying weekly label. We are glad to provide the "
 "companion manuscript to the editorial office on request.",
 "This manuscript is original, has not been published previously, and is not under consideration for "
 "publication elsewhere. All authors have read and approved the submitted version and agree to its "
 "submission. The study was approved by the Institutional Review Board of China Medical University Hospital "
 "(IRB No.: CRREC-108-001), and all participants provided written informed consent. The authors declare no "
 "conflict of interest.",
 "Thank you for considering our work; we look forward to your editorial assessment. We would be pleased to "
 "suggest qualified reviewers upon request [optional: list 2–3 names, affiliations and e-mails here].",
]
P2 = [
 "Please find enclosed our manuscript entitled “Interpretable Personalized Prediction of "
 "Smartphone-Addiction Risk from App Usage Time and Frequency: A Multi-Signal Ensemble with Green-Learning "
 "Features and SHAP Analysis,” which we respectfully submit for consideration as a Full-length report "
 "in the " + JOURNAL + ".",
 "Building on evidence that app-usage time tracks smartphone-addiction status, we asked which passively "
 "logged signals best support personalized, within-subject prediction of a student's weekly risk state. "
 "From 24 weeks of Android telemetry for 115 undergraduates, we treated app-usage time, usage frequency, "
 "functional category and app identity as candidate signals and, through a controlled design search, "
 "identified an interpretable rank-averaging ensemble that attained ROC-AUC 0.966 and PR-AUC 0.900 "
 "(per-fold 0.966 ± 0.009 and 0.901 ± 0.014; 95% CI [0.955, 0.977] and [0.884, 0.918]) under "
 "personalized within-subject cross-validation. SHAP and ablation analyses showed that usage time and "
 "frequency are the dominant, complementary predictors.",
 "The study offers three contributions of interest to the journal's readers: it shows that usage frequency "
 "is not redundant with duration; it provides a transparent, SHAP-interpretable model whose risk flags can "
 "be explained to clinicians and users; and it reports performance honestly across evaluation protocols, "
 "disclosing that strong within-subject monitoring performance does not transfer to cold-start prediction "
 "for unseen individuals. We therefore position the method as personalized monitoring rather than one-shot "
 "screening.",
 "In the interest of full transparency, a companion manuscript (“Detecting High-Risk Smartphone "
 "Addiction among College Students via AI-Based App Usage Analysis During COVID-19”) drawing on the same "
 "cohort is being submitted separately. The two papers are complementary and non-overlapping: the companion "
 "paper introduces a time-only detection pipeline with a fixed per-participant label, whereas the present "
 "paper contributes a multi-signal, interpretable ensemble for personalized, within-subject prediction of a "
 "time-varying weekly label, with SHAP-based attribution and an ablation study. We are glad to provide the "
 "companion manuscript to the editorial office on request.",
 "This manuscript is original, has not been published previously, and is not under consideration for "
 "publication elsewhere. All authors have read and approved the submitted version and agree to its "
 "submission. The study was approved by the Institutional Review Board of China Medical University Hospital "
 "(IRB No.: CRREC-108-001), and all participants provided written informed consent. The authors declare no "
 "conflict of interest.",
 "Thank you for considering our work; we look forward to your editorial assessment. We would be pleased to "
 "suggest qualified reviewers upon request [optional: list 2–3 names, affiliations and e-mails here].",
]
cover_letter("Paper1_CoverLetter.docx", P1)
cover_letter("Paper2_CoverLetter.docx", P2)
print("done")
